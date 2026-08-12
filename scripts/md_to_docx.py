#!/usr/bin/env python3
"""Convierte MANUAL_ADMINISTRADOR.md a Word (.docx)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(cell, 'EDE9FE')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()


def add_rich_paragraph(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    if bold:
        run = p.add_run(strip_md_inline(text))
        run.bold = True
        return p
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
        else:
            p.add_run(part)
    return p


def strip_md_inline(text):
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines = []
    table_rows = None

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip() == '```':
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith('```'):
            in_code = True
            code_lines = []
            i += 1
            continue

        if line.strip() == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        if '|' in line and line.strip().startswith('|'):
            if table_rows is None:
                table_rows = []
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            if i >= len(lines) or not (lines[i].strip().startswith('|')):
                add_table(doc, table_rows)
                table_rows = None
            continue

        if table_rows is not None:
            add_table(doc, table_rows)
            table_rows = None

        if line.startswith('# '):
            h = doc.add_heading(strip_md_inline(line[2:].strip()), level=0)
            h.runs[0].font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
            i += 1
            continue

        if line.startswith('## '):
            doc.add_heading(strip_md_inline(line[3:].strip()), level=1)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(strip_md_inline(line[4:].strip()), level=2)
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if m:
            add_rich_paragraph(doc, m.group(2), style='List Number')
            i += 1
            continue

        if line.strip().startswith('- '):
            add_rich_paragraph(doc, line.strip()[2:], style='List Bullet')
            i += 1
            continue

        if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip('*'))
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_rich_paragraph(doc, line.strip())
        i += 1

    if table_rows:
        add_table(doc, table_rows)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f'Generado: {docx_path}')


if __name__ == '__main__':
    root = Path(__file__).resolve().parents[1]
    md = root / 'MANUAL_ADMINISTRADOR.md'
    out = root / 'MANUAL_ADMINISTRADOR.docx'
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    convert(md, out)
